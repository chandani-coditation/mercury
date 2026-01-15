#!/usr/bin/env python3
"""Ingest runbooks from DOCX files.

This script extracts text content from DOCX files in the runbooks/ folder
and ingests them using JSON schema-driven extraction and field mappings configuration.

Usage:
    python scripts/data/ingest_runbooks.py --dir runbooks
    python scripts/data/ingest_runbooks.py --file "runbooks/Runbook - Database Alerts.docx"
"""
import argparse
import sys
import uuid
from pathlib import Path

# Add project root to path BEFORE importing project modules
project_root = Path(__file__).parent.parent.parent
sys.path.insert(0, str(project_root))

# Now import project modules
from typing import Dict, Optional
from docx import Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import Table
from docx.text.paragraph import Paragraph
from ai_service.core import get_field_mappings_config, get_logger, setup_logging
from ingestion.models import IngestRunbook
import requests
import json


# Default ingestion service URL
INGESTION_SERVICE_URL = "http://localhost:8002"


def clean_text(text: str) -> str:
    """Clean text by removing formatting artifacts and special characters."""
    if not text:
        return ""

    # Remove common formatting artifacts
    text = text.replace("¶", "")  # Remove paragraph markers
    text = text.replace("", "")  # Remove zero-width spaces
    text = text.replace("\u200b", "")  # Remove zero-width spaces (Unicode)
    text = text.replace("\xa0", " ")  # Replace non-breaking spaces with regular spaces

    # Remove markdown-style heading markers at the start
    text = text.lstrip("#").strip()

    # Remove multiple consecutive spaces
    while "  " in text:
        text = text.replace("  ", " ")

    # Remove multiple consecutive newlines
    while "\n\n\n" in text:
        text = text.replace("\n\n\n", "\n\n")

    return text.strip()


def extract_text_from_docx(docx_path: Path) -> Dict[str, any]:
    """Extract structured content from DOCX file.

    Returns dict with: title, steps, commands, prerequisites, rollback_procedures, content
    """
    doc = Document(docx_path)

    # Extract title (usually first paragraph or document property)
    title = docx_path.stem  # Default to filename
    if doc.paragraphs:
        first_para = clean_text(doc.paragraphs[0].text)
        if first_para and len(first_para) < 200:  # Likely a title
            title = first_para

    # Clean the title
    title = clean_text(title)

    # Extract all text content
    full_content_parts = []
    steps = []
    commands = []
    prerequisites = []
    rollback_steps = []
    rollback_commands = []
    rollback_preconditions = []
    rollback_triggers = []

    current_section = None
    remediation_steps = []  # Store remediation steps separately

    for element in doc.element.body:
        if isinstance(element, CT_P):
            para = Paragraph(element, doc)
            text = clean_text(para.text)

            if not text:
                continue

            # Detect section headers (usually bold or all caps)
            # But don't treat remediation step lines as headers even if bold
            is_header = False
            # Check if this looks like a remediation step first (has colon separating problem: solution)
            is_remediation_step = (
                current_section == "remediation"
                and ":" in text
                and len(text.split(":")) == 2
                and len(text.split(":")[1].strip()) > 10  # Has meaningful action text
            )

            if not is_remediation_step:
                # Only check for header if it's not a remediation step
                for run in para.runs:
                    if run.bold or text.isupper():
                        is_header = True
                        break

            if is_header:
                text_lower = text.lower()
                # Detect remediation sections (generic, not hard-coded)
                if (
                    "remediation" in text_lower
                    or "suggest remediation" in text_lower
                    or "remedy" in text_lower
                    or ("resolution" in text_lower and "output" not in text_lower)
                    or "mitigation" in text_lower
                    or "corrective action" in text_lower
                    or ("fix" in text_lower and len(text) > 10)  # Avoid false positives
                    or ("solution" in text_lower and len(text) > 10)
                ):
                    current_section = "remediation"
                elif "step" in text_lower or "procedure" in text_lower:
                    # Check if it's rollback steps
                    if "rollback" in text_lower or "revert" in text_lower:
                        current_section = "rollback_steps"
                    else:
                        current_section = "steps"
                elif "command" in text_lower or "cmd" in text_lower:
                    # Check if it's rollback commands
                    if "rollback" in text_lower or "revert" in text_lower:
                        current_section = "rollback_commands"
                    else:
                        current_section = "commands"
                elif "prerequisite" in text_lower or "requirement" in text_lower:
                    # Check if it's rollback preconditions
                    if "rollback" in text_lower or "revert" in text_lower:
                        current_section = "rollback_preconditions"
                    else:
                        current_section = "prerequisites"
                elif (
                    "rollback" in text_lower
                    or "revert" in text_lower
                    or "undo" in text_lower
                    or "restore" in text_lower
                ):
                    # Generic rollback section
                    current_section = "rollback"
                elif (
                    "trigger" in text_lower
                    or "when to rollback" in text_lower
                    or "rollback criteria" in text_lower
                ):
                    current_section = "rollback_triggers"
                elif (
                    "validate" in text_lower
                    and "recovery" in text_lower
                    or "documentation" in text_lower
                    and "requirement" in text_lower
                    or "threshold" in text_lower
                    and "tuning" in text_lower
                ):
                    # Clear remediation section when we hit these sections
                    current_section = None
                # Don't reset current_section if we're in remediation and hit a non-section header
                # (remediation steps might be formatted as headers)
                elif current_section != "remediation":
                    current_section = None

                # Add header without markdown markers
                full_content_parts.append(f"\n{text.upper()}\n{'='*len(text)}\n")
            else:
                # Add to appropriate section
                if current_section == "remediation":
                    # Extract remediation steps (can be colon-separated like "Connection saturation: Kill idle sessions")
                    # Only add if it's not empty and looks like a remediation step
                    if text and len(text.strip()) > 5:
                        remediation_steps.append(text)
                        full_content_parts.append(f"  • {text}\n")
                elif current_section == "steps":
                    steps.append(text)
                    full_content_parts.append(f"  • {text}\n")
                elif current_section == "commands":
                    # Commands might be in code blocks or plain text
                    if (
                        text.startswith("$")
                        or text.startswith("#")
                        or "sudo" in text
                        or "kubectl" in text
                    ):
                        commands.append(text)
                    full_content_parts.append(f"  $ {text}\n")
                elif current_section == "prerequisites":
                    prerequisites.append(text)
                    full_content_parts.append(f"  • {text}\n")
                elif current_section == "rollback" or current_section == "rollback_steps":
                    rollback_steps.append(text)
                    full_content_parts.append(f"  • {text}\n")
                elif current_section == "rollback_commands":
                    rollback_commands.append(text)
                    full_content_parts.append(f"  $ {text}\n")
                elif current_section == "rollback_preconditions":
                    rollback_preconditions.append(text)
                    full_content_parts.append(f"  • {text}\n")
                elif current_section == "rollback_triggers":
                    rollback_triggers.append(text)
                    full_content_parts.append(f"  {text}\n")
                else:
                    full_content_parts.append(f"{text}\n")

        elif isinstance(element, CT_Tbl):
            # Extract text from tables
            table = Table(element, doc)
            for row in table.rows:
                row_text = " | ".join([clean_text(cell.text) for cell in row.cells])
                if row_text:
                    full_content_parts.append(f"{row_text}\n")

    full_content = "".join(full_content_parts)

    # Prioritize remediation steps over generic steps
    # If remediation steps exist, use them; otherwise use generic steps
    final_steps = remediation_steps if remediation_steps else steps

    # Build structured rollback procedures
    rollback_procedures = None
    if rollback_steps:
        rollback_procedures = {
            "steps": rollback_steps,
            "commands": rollback_commands if rollback_commands else None,
            "preconditions": rollback_preconditions if rollback_preconditions else None,
            "triggers": rollback_triggers if rollback_triggers else None,
        }
        # If no structured data, fall back to text
        if not rollback_commands and not rollback_preconditions and not rollback_triggers:
            rollback_procedures = "\n".join(rollback_steps)

    return {
        "title": title,
        "steps": final_steps,
        "commands": commands,
        "prerequisites": prerequisites,
        "rollback_procedures": rollback_procedures,
        "content": full_content,
        "remediation_steps": remediation_steps,  # Keep separate for metadata
    }


def map_docx_to_runbook(docx_path: Path, field_mappings: Dict) -> IngestRunbook:
    """Map DOCX content to IngestRunbook using field mappings configuration."""
    # Extract structured content
    extracted = extract_text_from_docx(docx_path)

    # Use field mappings to determine which fields to use
    mappings = field_mappings.get("field_mappings", {})

    # Extract service/component from filename using config-driven pattern matching
    service = None
    component = None

    # Try to extract from filename with config-driven pattern matching
    filename_clean = docx_path.stem.replace("Runbook -", "").replace("Runbook –", "").strip()
    filename_lower = filename_clean.lower()

    # Load runbook filename patterns from config
    try:
        project_root = Path(__file__).parent.parent.parent
        mapping_path = project_root / "config" / "service_component_mapping.json"
        if mapping_path.exists():
            with open(mapping_path, "r") as f:
                service_component_mapping = json.load(f)
            runbook_patterns_config = service_component_mapping.get("runbook_filename_patterns", {})
            patterns = runbook_patterns_config.get("patterns", [])
            default_service = runbook_patterns_config.get("default_service", "General")
            component_suffixes = runbook_patterns_config.get(
                "component_suffixes_to_remove", [" Alerts", " Alert"]
            )
        else:
            patterns = []
            default_service = "General"
            component_suffixes = [" Alerts", " Alert"]
    except Exception:
        patterns = []
        default_service = "General"
        component_suffixes = [" Alerts", " Alert"]

    # Try each pattern in order (first match wins)
    matched = False
    for pattern in patterns:
        keywords = pattern.get("keywords", [])
        # Check if any keyword matches in filename
        if any(keyword in filename_lower for keyword in keywords):
            service = pattern.get("service", default_service)
            component_extraction = pattern.get("component_extraction", "remove_keywords")

            if component_extraction == "use_full_filename":
                component = filename_clean
            elif component_extraction == "remove_keywords":
                # Remove matched keywords from filename to get component
                component = filename_clean
                for keyword in keywords:
                    # Remove keyword (case-insensitive)
                    import re

                    component = re.sub(
                        re.escape(keyword), "", component, flags=re.IGNORECASE
                    ).strip()
                component = " ".join(component.split())  # Normalize whitespace
            else:
                component = None

            matched = True
            break

    # Fallback: if no pattern matched, use default logic
    if not matched:
        filename_parts = filename_clean.split()
        if len(filename_parts) >= 2:
            service = filename_parts[0]
            component = " ".join(filename_parts[1:])
        else:
            service = filename_parts[0] if filename_parts else default_service
            component = None

    # Clean up component (remove trailing suffixes like "Alerts")
    if component:
        component = component.strip()
        for suffix in component_suffixes:
            if component.endswith(suffix):
                component = component.replace(suffix, "").strip()
                break
        # If component is empty after cleanup, set to None
        if not component:
            component = None

    # Build comprehensive tags
    runbook_id = str(uuid.uuid4())
    tags = {
        "type": "runbook",
        "runbook_id": runbook_id,
        "source_file": docx_path.name,
    }

    if service:
        tags["service"] = service
    if component:
        tags["component"] = component

    # Build metadata
    metadata = {
        "source": "docx",
        "file_path": str(docx_path),
        "has_steps": len(extracted["steps"]) > 0,
        "has_commands": len(extracted["commands"]) > 0,
        "has_rollback": extracted["rollback_procedures"] is not None,
        "rollback_structured": isinstance(
            extracted["rollback_procedures"], dict
        ),  # Track if rollback is structured
    }

    return IngestRunbook(
        title=extracted["title"],
        service=service,
        component=component,
        content=extracted["content"],
        steps=extracted["steps"] if extracted["steps"] else None,
        prerequisites=extracted["prerequisites"] if extracted["prerequisites"] else None,
        rollback_procedures=extracted["rollback_procedures"],
        tags=tags,
        metadata=metadata,
    )


def ingest_runbook(
    runbook: IngestRunbook, ingestion_url: str = INGESTION_SERVICE_URL
) -> tuple[bool, Optional[str]]:
    """Ingest a single runbook via the ingestion API.

    Returns:
        Tuple of (success: bool, document_id: Optional[str])
    """
    logger = get_logger(__name__)
    try:
        response = requests.post(
            f"{ingestion_url}/ingest/runbook",
            json=runbook.model_dump(mode="json", exclude_none=True),
            timeout=60,  # Longer timeout for larger files
        )
        response.raise_for_status()
        result = response.json()
        document_id = result.get("document_id")
        return True, document_id
    except Exception as e:
        logger.error(f"Failed to ingest runbook {runbook.title}: {str(e)}")
        return False, None


def ingest_docx_file(
    file_path: Path,
    field_mappings: Dict,
    ingestion_url: str,
    file_num: int = None,
    total_files: int = None,
) -> tuple[int, int]:
    """Ingest a single DOCX file."""
    logger = get_logger(__name__)
    file_info = f"[{file_num}/{total_files}] " if file_num and total_files else ""
    logger.info(f"{file_info}Processing: {file_path.name}")

    try:
        runbook = map_docx_to_runbook(file_path, field_mappings)
        success, document_id = ingest_runbook(runbook, ingestion_url)

        if success:
            logger.info(f"Successfully ingested: {runbook.title[:50]}")
            return 1, 0
        else:
            logger.error(f"Failed to ingest runbook")
            return 0, 1

    except Exception as e:
        logger.error(f"Error processing {file_path.name}: {str(e)}")
        return 0, 1


def main():
    # Setup logging first to ensure output is visible
    setup_logging(log_level="INFO", service_name="ingestion_script")
    logger = get_logger(__name__)

    parser = argparse.ArgumentParser(description="Ingest runbooks from DOCX files")
    parser.add_argument("--dir", type=str, help="Directory containing DOCX files")
    parser.add_argument("--file", type=str, help="Single DOCX file to ingest")
    parser.add_argument(
        "--ingestion-url",
        type=str,
        default=INGESTION_SERVICE_URL,
        help=f"Ingestion service URL (default: {INGESTION_SERVICE_URL})",
    )

    args = parser.parse_args()

    if not args.dir and not args.file:
        parser.error("Either --dir or --file must be provided")

    try:
        field_mappings_config = get_field_mappings_config()
        runbook_mappings = field_mappings_config.get("runbook_docx", {})
    except Exception as e:
        logger.error(f" Failed to load field mappings: {str(e)}")
        sys.exit(1)

    total_success = 0
    total_errors = 0

    if args.file:
        # Process single file
        file_path = Path(args.file)
        if not file_path.exists():
            logger.error(f" File not found: {file_path}")
            sys.exit(1)

        if not file_path.suffix.lower() == ".docx":
            logger.error(f" File is not a DOCX file: {file_path}")
            sys.exit(1)

        success, errors = ingest_docx_file(file_path, runbook_mappings, args.ingestion_url)
        total_success += success
        total_errors += errors

    else:
        # Process directory
        dir_path = Path(args.dir)
        if not dir_path.exists():
            logger.error(f" Directory not found: {dir_path}")
            sys.exit(1)

        docx_files = list(dir_path.glob("*.docx"))
        if not docx_files:
            logger.warning(f"No DOCX files found in {dir_path}")
            sys.exit(0)

        for idx, docx_file in enumerate(docx_files, start=1):
            success, errors = ingest_docx_file(
                docx_file, runbook_mappings, args.ingestion_url, idx, len(docx_files)
            )
            total_success += success
            total_errors += errors

    logger.info(f"Ingestion Summary: {total_success} successful, {total_errors} errors")

    if total_success > 0:
        try:
            from db.connection import get_db_connection_context

            with get_db_connection_context() as conn:
                cur = conn.cursor()

                # Count runbook documents
                cur.execute("SELECT COUNT(*) FROM documents WHERE doc_type = 'runbook';")
                doc_result = cur.fetchone()
                doc_count = doc_result["count"] if isinstance(doc_result, dict) else doc_result[0]

                # Count runbook steps
                cur.execute("SELECT COUNT(*) FROM runbook_steps;")
                step_result = cur.fetchone()
                step_count = (
                    step_result["count"] if isinstance(step_result, dict) else step_result[0]
                )

                # Count runbook steps with embeddings
                cur.execute("SELECT COUNT(*) FROM runbook_steps WHERE embedding IS NOT NULL;")
                embed_result = cur.fetchone()
                embed_count = (
                    embed_result["count"] if isinstance(embed_result, dict) else embed_result[0]
                )

                cur.close()

            if embed_count == step_count and step_count > 0:
                logger.info(f"Verification: {step_count} runbook steps with embeddings")
            elif embed_count < step_count:
                logger.warning(f"Warning: {step_count - embed_count} steps missing embeddings")
            else:
                logger.warning("Warning: No runbook steps found in database")

        except Exception as e:
            logger.warning(f"Could not verify embeddings: {str(e)}")

    if total_errors > 0:
        logger.error(f"Completed with {total_errors} error(s). Check logs for details.")
        sys.exit(1)


if __name__ == "__main__":
    main()
